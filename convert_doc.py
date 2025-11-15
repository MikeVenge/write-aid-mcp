import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_table_from_data(doc, data):
    """Helper function to add a table to the document."""
    if not data:
        return
    
    # Assuming all rows have the same number of columns as the first row
    num_cols = len(data[0])
    table = doc.add_table(rows=len(data), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            # Ensure cell_data is a string
            if cell_data is None:
                cell_data = ""
            elif not isinstance(cell_data, str):
                cell_data = str(cell_data)
                
            row.cells[j].text = cell_data
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)

def main():
    try:
        doc = Document()
        doc.add_heading('Adgorithmics株式会社', level=0)
        doc.add_heading('ARTICLES OF INCORPORATION', level=1)
        doc.add_paragraph('Source: T02-2608391-1.0-Documants for incorporation (1) (1).pdf')
        doc.add_paragraph()

        # --- PAGE 1-2 ---
        doc.add_paragraph('[Original Copy Text]', 'Heading 3')
        [cite_start]doc.add_paragraph('原本 [cite: 5]')
        [cite_start]doc.add_paragraph('Adgorithmics株式会社 [cite: 6]')
        [cite_start]doc.add_paragraph('定 款 [cite: 7, 8]')
        
        [cite_start]doc.add_heading('ARTICLES OF INCORPORATION OF ADGORITHMICS KK [cite: 16, 17, 18]', level=2)
        
        [cite_start]doc.add_heading('第1章 総則 [cite: 19]', level=3)
        [cite_start]doc.add_heading('CHAPTER I. GENERAL PROVISIONS [cite: 20]', level=3)
        
        [cite_start]doc.add_paragraph('[Notary Stamp] [cite: 21]', 'Body Text')
        
        [cite_start]doc.add_paragraph().add_run('(商号) [cite: 22]').bold = True
        [cite_start]doc.add_paragraph('第1条 当会社は、Adgorithmics株式会社と称し、英文では、Adgorithmics KKと表示する。 [cite: 23]')
        
        [cite_start]doc.add_paragraph().add_run('(Company Name) [cite: 24]').bold = True
        [cite_start]doc.add_paragraph('Article 1. [cite: 25][cite_start]\nThe name of the Company shall be Adgorithmics Kabushiki Kaisha, and shall be indicated as Adgorithmics KK in English. [cite: 26, 27]')
        
        [cite_start]doc.add_paragraph().add_run('(目的) [cite: 28]').bold = True
        [cite_start]doc.add_paragraph('第2条 当会社は、次の事業を営むことを目的とする。 [cite: 29]')
        doc.add_paragraph(
            [cite_start]'(1) ソフトウェアの調査、開発 [cite: 30]\n'
            [cite_start]'(2) ソフトウェアの販売、マーケティング [cite: 31]\n'
            [cite_start]'(3) ハードウェアの調査、開発 [cite: 32]\n'
            [cite_start]'(4) ハードウェアの販売、マーケティング [cite: 33]\n'
            [cite_start]'(5) セミナー、会議の開催 [cite: 34]\n'
            [cite_start]'(6) デジタルメディア、ソーシャル・ネットワーキング・サービス、テレビ、雑誌、新聞、看板その他あらゆる方式の広告の代理店業 [cite: 35]\n'
            [cite_start]'(7) デジタルメディア、ソーシャル・ネットワーキング・サービス、テレビ、雑誌、新聞、看板その他あらゆる方式の広告のコンテンツ(ビデオ、写真、テキスト、図形、音声、ポスター、画像など)のデザイン、制作 [cite: 36, 37]\n'
            [cite_start]'(8) 海外のソフトウェア、ハードウェア製品の日本への輸入販売 [cite: 38]\n'
            [cite_start]'(9) 国内のソフトウェア、ハードウェア製品の海外への輸出販売 [cite: 39]\n'
            [cite_start]'(10) 各種マーケティング、販売業務 [cite: 40]\n'
            [cite_start]'(11) デジタルメディア、ソーシャル・ネットワーキング・サービス、テレビ、雑誌、新聞、看板その他あらゆる方式の広告のためのプラットフォーム、システム、プログラミング処理の調査、開発 [cite: 41]\n'
            [cite_start]'(12) 前各号に付帯関連する一切の業務 [cite: 43]')
        
        # --- PAGE 3 ---
        [cite_start]doc.add_paragraph().add_run('(Purpose) [cite: 44]').bold = True
        [cite_start]doc.add_paragraph('Article 2. [cite: 45][cite_start]\nThe purpose of the Company shall be to engage in the following business activities: [cite: 46]')
        doc.add_paragraph(
            [cite_start]'(1) Research and development of software; [cite: 47]\n'
            [cite_start]'(2) Sales and marketing of software; [cite: 48]\n'
            [cite_start]'(3) Research and development of hardware; [cite: 49]\n'
            [cite_start]'(4) Sales and marketing of hardware; [cite: 50]\n'
            [cite_start]'(5) Holding seminars and meetings; [cite: 51]\n'
            [cite_start]'(6) Advertising agency business through digital media, social networking services, TV, magazines, newspapers, billboards or any other means; [cite: 52]\n'
            [cite_start]'(7) Design and production of content in the form of videos, photographs, text, figures, sound, posters, graphics and so on for advertising through digital media, social networking services, TV, magazines, newspapers, billboards or any other means; [cite: 53]\n'
            [cite_start]'(8) Import and sales of software and hardware products made in foreign countries; [cite: 54]\n'
            [cite_start]'(9) Export and sales of software and hardware products made in Japan; [cite: 55]\n'
            [cite_start]'(10) All kinds of marketing and sales business; [cite: 56]\n'
            [cite_start]'(11) Research and development of platforms, systems, programming processes for advertising through digital media, social networking services, TV, magazines, newspapers, billboards or any other means; and [cite: 57, 58]\n'
            [cite_start]'(12) All other businesses related to or connected with the activities described in each item above. [cite: 59]')
            
        [cite_start]doc.add_paragraph().add_run('(本店の所在地) [cite: 60]').bold = True
        [cite_start]doc.add_paragraph('第3条 当会社は、本店を東京都港区に置く。 [cite: 61]')
        
        [cite_start]doc.add_paragraph().add_run('(Location of Head Office) [cite: 62]').bold = True
        [cite_start]doc.add_paragraph('Article 3. [cite: 63][cite_start]\nThe head office of the Company shall be located at Minato-ku, Tokyo, Japan. [cite: 64]')
        
        [cite_start]doc.add_paragraph().add_run('(公告方法) [cite: 65]').bold = True
        [cite_start]doc.add_paragraph('第4条 当会社の公告は、官報に掲載して行う。 [cite: 66]')
        
        [cite_start]doc.add_paragraph().add_run('(Method of Public Notices) [cite: 67]').bold = True
        [cite_start]doc.add_paragraph('Article 4. [cite: 68][cite_start]\nPublic notices of the Company shall be made in the Official Gazette (kampo). [cite: 69]')

        # --- PAGE 4 ---
        [cite_start]doc.add_heading('第2章 株式 [cite: 71]', level=3)
        [cite_start]doc.add_heading('CHAPTER II. SHARES [cite: 72]', level=3)
        [cite_start]doc.add_paragraph('[Notary Stamp] [cite: 73]', 'Body Text')
        
        [cite_start]doc.add_paragraph().add_run('(発行可能株式総数) [cite: 74]').bold = True
        [cite_start]doc.add_paragraph('第5条 当会社の発行可能株式総数は、400株とする。 [cite: 75]')
        [cite_start]doc.add_paragraph().add_run('(Total Number of Issuable Shares) [cite: 76]').bold = True
        [cite_start]doc.add_paragraph('Article 5. [cite: 77][cite_start]\nThe total number of issuable shares of the Company shall be four hundred (400) shares. [cite: 78]')
        
        [cite_start]doc.add_paragraph().add_run('(株券の不発行) [cite: 79]').bold = True
        [cite_start]doc.add_paragraph('第6条 当会社は、株式に係る株券を発行しない。 [cite: 80]')
        [cite_start]doc.add_paragraph().add_run('(Non-issuance of Share Certificates) [cite: 81]').bold = True
        [cite_start]doc.add_paragraph('Article 6. [cite: 82][cite_start]\nThe Company shall not issue any share certificates for its shares. [cite: 83]')
        
        [cite_start]doc.add_paragraph().add_run('(株式の譲渡制限) [cite: 84]').bold = True
        [cite_start]doc.add_paragraph('第7条 当会社の株式を譲渡により取得することについて、株主または株式取得者は株主総会の承認を受けなければならない。 [cite: 85]')
        [cite_start]doc.add_paragraph().add_run('(Restriction on the Transfer of Shares) [cite: 86]').bold = True
        [cite_start]doc.add_paragraph('Article 7. [cite: 87][cite_start]\nA shareholder or an acquirer of shares of the Company shall obtain the approval of the General Meeting of Shareholders concerning the acquisition of shares by transfer. [cite: 88, 89]')

        [cite_start]doc.add_paragraph().add_run('(株主名簿記載事項の記載の請求) [cite: 90]').bold = True
        [cite_start]doc.add_paragraph('第8条 当会社の株式取得者が株主名簿記載事項を株主名簿に記載または記録することを請求するには、株式取得者とその取得した株式の株主として株主名簿に記載され、もしくは記録された者またはその相続人その他の一般承継人が当会社所定の書式による請求書に署名または記名押印し、共同して請求しなければならない。ただし、法令に別段の定めがある場合には、株式取得者が単独で請求することができる。 [cite: 91]')
        
        # --- PAGE 5 ---
        [cite_start]doc.add_paragraph().add_run('(Request for Stating the Items Listed in the Shareholder Registry) [cite: 92]').bold = True
        [cite_start]doc.add_paragraph('Article 8. [cite: 92][cite_start]\nIn requesting that the Company list or record in the Shareholder Registry the items listed in the Shareholder Registry, the acquirer of the shares shall, jointly with the person listed or recorded in the Shareholder Registry as the shareholder of the shares acquired, or this person\'s heir, or any other general successor, submit to the Company a written application in the form prescribed by the Company, which bears the signatures of, or the names and seals of the acquirer and the foregoing parties (as applicable); [cite: 93, 94, 95][cite_start]\nprovided, however, that, if otherwise provided for by any applicable laws or regulations, the acquirer of the shares may individually request the Company state or record those items listed in the Shareholder Registry. [cite: 96]')
        
        [cite_start]doc.add_paragraph().add_run('(株主の住所等の届出等) [cite: 97]').bold = True
        [cite_start]doc.add_paragraph('第9条 当会社の株主、登録株式質権者またはその法定代理人もしくは代表者は、当会社所定の書式により、その氏名または名称および住所ならびに印鑑を当会社に届け出なければならない。 [cite: 98]')
        [cite_start]doc.add_paragraph('2 前項の届出事項に変更を生じた場合も、同様とする。 [cite: 99]')
        [cite_start]doc.add_paragraph('3 当会社に提出する書類には、本条により届け出た印鑑を用いなければならない。 [cite: 100]')
        
        [cite_start]doc.add_paragraph().add_run('(Notification, etc. of Address and Other Matters Pertaining to Shareholders) [cite: 101]').bold = True
        [cite_start]doc.add_paragraph('Article 9. [cite: 102]\n1. [cite_start]Shareholders, registered pledgees of shares, and their statutory agents or representatives shall notify the Company of their names or company names, addresses and submit any seal impressions in the form prescribed by the Company. [cite: 103, 104]')
        [cite_start]doc.add_paragraph('2. Any changes to the matters to be notified or submitted, as stipulated in the preceding paragraph, shall be notified or submitted to the Company using the same method. [cite: 105, 106]')
        [cite_start]doc.add_paragraph('3. Seal impressions, identical to the seal impressions submitted to the Company under this Article, shall be used on any documents to be submitted to the Company. [cite: 107, 108]')

        # --- PAGE 6 ---
        [cite_start]doc.add_heading('第3章 株主総会 [cite: 109]', level=3)
        [cite_start]doc.add_heading('CHAPTER III. GENERAL MEETING OF SHAREHOLDERS [cite: 110]', level=3)
        
        [cite_start]doc.add_paragraph().add_run('(株主総会の招集) [cite: 111]').bold = True
        [cite_start]doc.add_paragraph('第10条 当会社の定時株主総会は、毎事業年度末日の翌日から3か月以内に招集し、臨時株主総会は、必要あるときに随時これを招集する。 [cite: 112]')
        [cite_start]doc.add_paragraph().add_run('(Convocation of General Meeting of Shareholders) [cite: 113]').bold = True
        [cite_start]doc.add_paragraph('Article 10. [cite: 114][cite_start]\nThe Company\'s ordinary General Meeting of Shareholders shall be convened within three (3) months from the day after the account closing date of each business year, and an extraordinary General Meeting of Shareholders shall be convened at any time as necessary. [cite: 115]')
        
        [cite_start]doc.add_paragraph().add_run('(定時株主総会の基準日) [cite: 118]').bold = True
        [cite_start]doc.add_paragraph('第11条 当会社の定時株主総会の議決権の基準日は、毎事業年度末日とする。 [cite: 119]')
        [cite_start]doc.add_paragraph().add_run('(Record Date for General Meeting of Shareholders) [cite: 120]').bold = True
        [cite_start]doc.add_paragraph('Article 11. [cite: 121][cite_start]\nThe record date concerning voting rights at the Company\'s ordinary General Meeting of Shareholders is the last day of each business year. [cite: 122]')

        [cite_start]doc.add_paragraph().add_run('(招集権者および議長) [cite: 123]').bold = True
        [cite_start]doc.add_paragraph('第12条 株主総会は、代表取締役がこれを招集し、議長となる。 [cite: 124]')
        [cite_start]doc.add_paragraph().add_run('(Convener and Chairperson) [cite: 125]').bold = True
        [cite_start]doc.add_paragraph('Article 12. [cite: 126][cite_start]\nThe Representative Director of the Company shall convene the General Meeting of Shareholders and shall act as the chairperson of the meeting. [cite: 127]')

        [cite_start]doc.add_paragraph().add_run('(招集手続) [cite: 128]').bold = True
        [cite_start]doc.add_paragraph('第13条 株主総会の招集通知は、会日の3日前までに議決権を行使することができる株主に対して発する。なお、招集通知は、書面ですることを要しない。 [cite: 129]')
        [cite_start]doc.add_paragraph('2 議決権を行使することができる株主の全員の同意があるときは、法令に別段の定めがある場合を除き、招集の手続を経ないで株主総会を開催することができる。 [cite: 130]')
        [cite_start]doc.add_paragraph().add_run('(Convocation Procedure) [cite: 131]').bold = True
        [cite_start]doc.add_paragraph('Article 13. [cite: 132]\n1. Notice of convocation of a General Meeting of Shareholders shall be dispatched to shareholders, entitled to exercise their votes at the meeting, at least three (3) days before the date of the meeting. [cite_start]Convocation notice need not be given in writing. [cite: 133, 134]')
        [cite_start]doc.add_paragraph('2. Unless otherwise provided for by any applicable laws or regulations, a General Meeting of Shareholders may be held without undertaking the convocation procedure if the consent of all shareholders entitled to exercise their votes is obtained. [cite: 135, 136]')

        # --- PAGE 7 ---
        [cite_start]doc.add_paragraph().add_run('(決議の方法) [cite: 137]').bold = True
        [cite_start]doc.add_paragraph('第14条 株主総会の決議は、法令または本定款に別段の定めがある場合を除き、議決権を行使することができる株主の議決権の過半数を有する株主が出席し、出席した当該株主の議決権の過半数をもって行う。 [cite: 138]')
        [cite_start]doc.add_paragraph().add_run('(Method of Resolution) [cite: 139]').bold = True
        [cite_start]doc.add_paragraph('Article 14. [cite: 140][cite_start]\nUnless otherwise provided for by any applicable laws or regulations or these Articles of Incorporation, a resolution of a General Meeting of Shareholders shall be adopted by a majority of the votes of the shareholders present at the meeting where the shareholders holding a majority of the votes of the shareholders who are entitled to exercise their votes are present. [cite: 142]')

        [cite_start]doc.add_heading('第4章 取締役 [cite: 143]', level=3)
        [cite_start]doc.add_heading('CHAPTER IV. DIRECTORS [cite: 144, 148]', level=3)

        [cite_start]doc.add_paragraph().add_run('(員数) [cite: 145]').bold = True
        [cite_start]doc.add_paragraph('第15条 [cite: 146][cite_start]\n当会社の取締役は、1名以上とする。 [cite: 147]')
        [cite_start]doc.add_paragraph().add_run('(Number of Directors) [cite: 149]').bold = True
        [cite_start]doc.add_paragraph('Article 15. [cite: 150][cite_start]\nThe Company shall have one (1) or more Directors. [cite: 151]')

        [cite_start]doc.add_paragraph().add_run('(選任方法) [cite: 152]').bold = True
        [cite_start]doc.add_paragraph('第16条 取締役は、株主総会において選任する。 [cite: 153]')
        [cite_start]doc.add_paragraph('2 取締役の選任決議は、累積投票によらないものとする。 [cite: 154]')
        [cite_start]doc.add_paragraph().add_run('(Method of Election) [cite: 155]').bold = True
        [cite_start]doc.add_paragraph('Article 16. [cite: 156]\n1. [cite_start]The Directors shall be elected at a General Meeting of Shareholders. [cite: 157]')
        [cite_start]doc.add_paragraph('2. A resolution for the election of the Directors cannot be adopted based on cumulative votes. [cite: 158]')

        # --- PAGE 8 ---
        [cite_start]doc.add_paragraph().add_run('(任期) [cite: 159]').bold = True
        [cite_start]doc.add_paragraph('第17条 取締役の任期は、選任後10年以内に終了する事業年度のうち最終のものに関する定時株主総会の終結の時までとする。 [cite: 160]')
        [cite_start]doc.add_paragraph('2 増員または補欠として選任された取締役の任期は、他の在任取締役または前任者の任期の満了する時までとする。 [cite: 161]')
        [cite_start]doc.add_paragraph().add_run('(Term of Office) [cite: 162]').bold = True
        [cite_start]doc.add_paragraph('Article 17. [cite: 163]\n1. [cite_start]The term of office of a Director shall continue until the conclusion of the ordinary General Meeting of Shareholders for the last business year which ends within ten (10) year period from the time of his or her election. [cite: 164]')
        [cite_start]doc.add_paragraph('2. The term of office of a Director who has been elected to increase the number of Directors or to fill the vacancy of a Director who resigned prior to the expiration of his or her term of office shall be the same as the remaining term of office of the other Directors currently in office or of the predecessor. [cite: 165, 166, 168, 169]')

        [cite_start]doc.add_paragraph().add_run('(業務の執行) [cite: 170]').bold = True
        [cite_start]doc.add_paragraph('第18条 取締役が2名以上ある場合には、当会社の業務は、取締役の過半数をもって決定する。 [cite: 171]')
        [cite_start]doc.add_paragraph().add_run('(Execution of Operations) [cite: 172]').bold = True
        [cite_start]doc.add_paragraph('Article 18. [cite: 173][cite_start]\nIn cases where there are two (2) or more Directors, the operations of the Company shall be decided by a majority of Directors. [cite: 174]')

        [cite_start]doc.add_paragraph().add_run('(代表取締役) [cite: 175]').bold = True
        [cite_start]doc.add_paragraph('第19条 取締役が2名以上ある場合には、株主総会の決議により、取締役の中から代表取締役1名を定める。 [cite: 176]')
        [cite_start]doc.add_paragraph().add_run('(Representative) [cite: 177]').bold = True
        [cite_start]doc.add_paragraph('Article 19. [cite: 178][cite_start]\nIn cases where there are two (2) or more Directors, subject to a resolution of a General Meeting of Shareholders, the Company shall elect one (1) Representative Director out of the Directors. [cite: 179]')

        # --- PAGE 9 ---
        [cite_start]doc.add_heading('第5章 計算 [cite: 180]', level=3)
        [cite_start]doc.add_heading('CHAPTER V. ACCOUNTING [cite: 181, 182]', level=3)
        
        [cite_start]doc.add_paragraph().add_run('(事業年度) [cite: 183]').bold = True
        [cite_start]doc.add_paragraph('第20条 当会社の事業年度は、毎年1月1日から同年12月31日までの1年とする。 [cite: 184]')
        [cite_start]doc.add_paragraph().add_run('(Business Year) [cite: 185]').bold = True
        [cite_start]doc.add_paragraph('Article 20. [cite: 186][cite_start]\nThe business year of the Company shall commence on January 1 of each year and end on December 31 of the same year. [cite: 187]')

        [cite_start]doc.add_paragraph().add_run('(剰余金の配当) [cite: 188]').bold = True
        [cite_start]doc.add_paragraph('第21条 当会社は、株主総会の決議によって、剰余金の配当を行う。 [cite: 189]')
        [cite_start]doc.add_paragraph('2 当会社の期末配当の基準日は、毎事業年度末日とする。 [cite: 190]')
        [cite_start]doc.add_paragraph().add_run('(Distribution of Surplus) [cite: 191]').bold = True
        [cite_start]doc.add_paragraph('Article 21. [cite: 193]\n1. [cite_start]The Company shall distribute dividends of its surplus based on a resolution of a General Meeting of Shareholders. [cite: 194]')
        [cite_start]doc.add_paragraph('2. The record date for the Company\'s year-end dividends shall be the last day of each business year. [cite: 195, 196]')

        [cite_start]doc.add_paragraph().add_run('(配当財産の除斥期間) [cite: 197]').bold = True
        [cite_start]doc.add_paragraph('第22条 配当財産がその交付開始の日から満3年を経過してもなお受領されないときは、当会社はその交付義務を免れる。 [cite: 198]')
        [cite_start]doc.add_paragraph().add_run('(Limitation of Period of Asset to be Distributed) [cite: 199]').bold = True
        [cite_start]doc.add_paragraph('Article 22. [cite: 200][cite_start]\nIf any shareholder fails to collect any dividend property within three (3) full years from the date on which the asset to be distributed is provided, the Company shall be exempted from its obligation to provide the asset to be distributed. [cite: 201]')

        # --- PAGE 10 ---
        [cite_start]doc.add_heading('第6章 附則 [cite: 202]', level=3)
        [cite_start]doc.add_heading('CHAPTER VI. SUPPLEMENTARY PROVISIONS [cite: 203, 204]', level=3)

        [cite_start]doc.add_paragraph().add_run('(設立に際して出資される財産の価額) [cite: 205]').bold = True
        [cite_start]doc.add_paragraph('第23条 当会社の設立に際して出資される財産の価額は金10万円とする。 [cite: 206]')
        [cite_start]doc.add_paragraph().add_run('(Value of Property to Be Contributed at Incorporation) [cite: 207]').bold = True
        [cite_start]doc.add_paragraph('Article 23. [cite: 208][cite_start]\nThe total value of the property to be contributed at the incorporation of the Company is 100,000 yen. [cite: 209, 210]')

        [cite_start]doc.add_paragraph().add_run('(発起人の氏名または名称および住所) [cite: 211]').bold = True
        [cite_start]doc.add_paragraph('第24条 当会社の発起人の氏名または名称および住所は、次のとおりとする。 [cite: 212]')
        [cite_start]doc.add_paragraph('東京都港区六本木六丁目12番2-1110号 [cite: 213][cite_start]\nキム・マイケル・サン・エン [cite: 214]')
        [cite_start]doc.add_paragraph().add_run('(Name or Company Name, and Address of Incorporator) [cite: 215]').bold = True
        [cite_start]doc.add_paragraph('Article 24. [cite: 216][cite_start]\nThe name or the entity name and address of the Incorporator of the Company are as follows: [cite: 217, 218]')
        [cite_start]doc.add_paragraph('Kim Michael Sang Eun of 6-12-2-1110 Roppongi, Minato-ku, Tokyo, Japan [cite: 219]')

        [cite_start]doc.add_paragraph().add_run('(最初の事業年度) [cite: 221]').bold = True
        [cite_start]doc.add_paragraph('第25条 当会社の最初の事業年度は、第20条の規定にかかわらず、当会社成立の日から平成28年12月31日までとする。 [cite: 222]')
        [cite_start]doc.add_paragraph().add_run('(Initial Business Year) [cite: 223]').bold = True
        [cite_start]doc.add_paragraph('Article 25. [cite: 224][cite_start]\nNotwithstanding Article 20, the initial business year of the Company shall commence on the date of its incorporation and end on December 31, 2016. [cite: 225, 226]')
        
        [cite_start]doc.add_paragraph().add_run('(設立時取締役) [cite: 227]').bold = True
        [cite_start]doc.add_paragraph('第26条 当会社の設立時取締役は、次のとおりとする。 [cite: 228][cite_start]\n設立時取締役 キム・マイケル・サン・エン [cite: 229]')
        [cite_start]doc.add_paragraph().add_run('(Director at Incorporation) [cite: 230]').bold = True
        [cite_start]doc.add_paragraph('Article 26. [cite: 231][cite_start]\nDirector at Incorporation shall be as follows: [cite: 232][cite_start]\nDirector at incorporation: Kim Michael Sang Eun [cite: 233, 234]')

        [cite_start]doc.add_paragraph('以上、Adgorithmics株式会社設立のため、発起人は、この定款を作成し、署名または記名押印する。 [cite: 235]')
        [cite_start]doc.add_paragraph('IN WITNESS WHEREOF, the Incorporator has affixed his or her signature or name and seal onto these Articles of Incorporation on December 2, 2016. [cite: 236]')
        [cite_start]doc.add_paragraph('平成28年12月7日 [cite: 237]')
        [cite_start]doc.add_paragraph('Adgorithmics株式会社 [cite: 238][cite_start]\n発起人 キム・マイケル・サン・エン [cite: 239]')
        [cite_start]doc.add_paragraph('Adgorithmics KK [cite: 240][cite_start]\nIncorporator: Kim Michael Sang Eun [cite: 241]')
        [cite_start]doc.add_paragraph('[KIM] [cite: 242][cite_start]\n[Personal seal] [cite: 243, 245]')

        # --- PAGE 11 ---
        doc.add_page_break()
        doc.add_heading('Notarization Page', level=2)
        [cite_start]doc.add_paragraph('平成28年登簿第130号 [cite: 246]')
        [cite_start]doc.add_paragraph('この定款における発起人 KIM MICHAEL SANG EUN (キム・マイケル・サン・エン)の代理人松並祐未は、本職の面前で発起人が自己の記名捺印を自認する旨陳述した。 [cite: 246]')
        [cite_start]doc.add_paragraph('よって、これを認証する。 [cite: 246]')
        [cite_start]doc.add_paragraph('平成28年12月8日,本職役場において [cite: 246]')
        [cite_start]doc.add_paragraph('東京都千代田区内幸町2丁目2番2号 [cite: 246]')
        [cite_start]doc.add_paragraph('東京法務局所属 公証人 大野 UB 役場 [cite: 246, 247]')
        
        # --- PAGE 12 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('発起人決定書 [cite: 249]', level=1)
        [cite_start]doc.add_heading('STATEMENT OF INCORPORATOR [cite: 250]', level=1)
        [cite_start]doc.add_paragraph('平成28年12月早日、Adgorithmics株式会社(以下「当会社」という。)の発起人は、東京都港区北青山二丁目9番5号スタジアムプレイス青山8階当会社創立事務所において、下記事項を定めた。2 [cite: 251]')
        [cite_start]doc.add_paragraph('On December, 2016, the Incorporator of Adgorithmics KK (the "Company"), determined the following matters at 8th Floor STUDIUM PLACE AOYAMA, 2-9-5 Kita-aoyama, Minato-ku, Tokyo, Japan, the organization office of the Company: [cite: 252]')
        [cite_start]doc.add_paragraph('記 [cite: 253][cite_start]\nNote [cite: 254]')
        
        [cite_start]doc.add_paragraph('1. 商号 [cite: 255][cite_start]\nAdgorithmics株式会社 [cite: 256][cite_start]\n(英文表記 Adgorithmics KK) [cite: 258]')
        [cite_start]doc.add_paragraph('Company name: [cite: 257][cite_start]\nAdgorithmics Kabushiki Kaisha [cite: 258][cite_start]\n(Adgorithmics KK in English) [cite: 259]')
        
        [cite_start]doc.add_paragraph('2. 目的 [cite: 260][cite_start]\nPurpose: [cite: 261]')
        [cite_start]doc.add_paragraph('当会社は、次の事業を営むことを目的とする。 [cite: 262]')
        doc.add_paragraph(
            [cite_start]'(1) ソフトウェアの調査、開発 [cite: 263]\n'
            [cite_start]'(2) ソフトウェアの販売、マーケティング [cite: 264]\n'
            [cite_start]'(3) ハードウェアの調査、開発 [cite: 265]\n'
            [cite_start]'(4) ハードウェアの販売、マーケティング [cite: 266]\n'
            [cite_start]'(5) セミナー、会議の開催 [cite: 267]\n'
            [cite_start]'(6) デジタルメディア、ソーシャル・ネットワーキング・サービス、テレビ、雑誌、新聞、看板その他あらゆる方式の広告の代理店業 [cite: 268, 269]\n'
            [cite_start]'(7) デジタルメディア、ソーシャル・ネットワーキング・サービス、テレビ、雑誌、新聞、看板その他あらゆる方式の広告のコンテンツ(ビデオ、写真、テキスト、図形、音声、ポスター、画像など)のデザイン、制作 [cite: 270]\n'
            [cite_start]'(8) 海外のソフトウェア、ハードウェア製品の日本への輸入販売 [cite: 271]\n'
            [cite_start]'(9) 国内のソフトウェア、ハードウェア製品の海外への輸出販売 [cite: 272]\n'
            [cite_start]'(10) 各種マーケティング、販売業務 [cite: 273]\n'
            [cite_start]'(11) デジタルメディア、ソーシャル・ネットワーキング・サービス、テレビ、雑誌、新聞、看板その他あらゆる方式の広告のためのプラットフォーム、システム、プログラミング処理の調査、開発 [cite: 274, 275]\n'
            [cite_start]'(12) 前各号に付帯関連する一切の業務 [cite: 276]')
        [cite_start]doc.add_paragraph('The purpose of the Company shall be to engage in the following business activities: [cite: 277]')
        doc.add_paragraph(
            [cite_start]'(1) Research and development of software; [cite: 278]\n'
            [cite_start]'(2) Sales and marketing of software; [cite: 279]\n'
            [cite_start]'(3) Research and development of hardware; [cite: 280]\n'
            [cite_start]'(4) Sales and marketing of hardware; [cite: 281]\n'
            [cite_start]'(5) Holding seminars and meetings; [cite: 282]\n'
            [cite_start]'(6) Advertising agency business through digital media, social networking services, TV, magazines, newspapers, billboards or any other means; [cite: 283, 284]\n'
            [cite_start]'(7) Design and production of content in the form of videos, photographs, text, figures, sound, posters, graphics and so on for advertising through digital media, social networking services, TV, magazines, newspapers, billboards or any other means; [cite: 285]\n'
            [cite_start]'(8) Import and sales of software and hardware products made in foreign countries; [cite: 286]\n'
            [cite_start]'(9) Export and sales of software and hardware products made in Japan; [cite: 287]')
            
        # --- PAGE 13 ---
        doc.add_paragraph(
            [cite_start]'(10) All kinds of marketing and sales business; [cite: 288]\n'
            [cite_start]'(11) Research and development of platforms, systems, programming processes for advertising through digital media, social networking services, TV, magazines, newspapers, billboards or any other means; and [cite: 289, 290]\n'
            [cite_start]'(12) All other businesses related to or connected with the activities described in each item above. [cite: 291]')

        [cite_start]doc.add_paragraph('3. 本店の所在地 [cite: 293][cite_start]\n東京都港区 [cite: 292]')
        [cite_start]doc.add_paragraph('Location of the head office: [cite: 294][cite_start]\nMinato-ku, Tokyo, Japan [cite: 295]')
        
        [cite_start]doc.add_paragraph('4. 発行可能株式総数 [cite: 296][cite_start]\n400株 [cite: 297]')
        [cite_start]doc.add_paragraph('Total number of issuable shares: [cite: 298][cite_start]\n400 shares [cite: 299]')
        
        [cite_start]doc.add_paragraph('5. 株式の譲渡制限 [cite: 300][cite_start]\n当会社の株式を譲渡により取得することについて、株主または株式取得者は株主総会の承認を受けなければならない。 [cite: 301, 302]')
        [cite_start]doc.add_paragraph('Restriction on the transfer of shares: [cite: 303][cite_start]\nA shareholder or an acquirer of shares of the Company shall obtain the approval of the General Meeting of Shareholders concerning the acquisition of shares by transfer. [cite: 304, 305]')
        
        [cite_start]doc.add_paragraph('6. 設立に際して出資される財産の価額 [cite: 306][cite_start]\n金100,000円 [cite: 307]')
        [cite_start]doc.add_paragraph('Total value of the assets to be contributed at the incorporation: [cite: 308][cite_start]\n100,000 yen [cite: 308]')

        [cite_start]doc.add_paragraph('7. 発起人の氏名および住所ならびに発起人が割当てを受ける設立時発行株式の数および設立時発行株式と引換えに払い込む金銭の額 [cite: 309]')
        [cite_start]doc.add_paragraph('Name and address of the Incorporator, the number of issued share at the incorporation allocated to the Incorporator, and the subscription payment for this share: [cite: 310, 311]')
        [cite_start]doc.add_paragraph('住所: 東京都港区六本木六丁目12番2-1110号 [cite: 312][cite_start]\n氏名: キム・マイケル・サン・エン [cite: 313][cite_start]\n割当てを受ける株式数: 1株 [cite: 314][cite_start]\n払込金額:100,000円 [cite: 315]')
        [cite_start]doc.add_paragraph('Address: 6-12-2-1110 Roppongi, Minato-ku, Tokyo, Japan [cite: 316][cite_start]\nName: Kim Michael Sang Eun [cite: 317][cite_start]\nNumber of shares to be allotted: 1 share [cite: 318][cite_start]\nAmount of payment therefor: 100,000 yen [cite: 319]')
        
        [cite_start]doc.add_paragraph('8. 成立後の資本金および資本準備金の額 [cite: 320][cite_start]\nMatters regarding the amount of the stated capital and capital reserve: [cite: 327]')
        [cite_start]doc.add_paragraph('資本金の額 [cite: 321][cite_start]\n金100,000円 [cite: 325]')
        [cite_start]doc.add_paragraph('資本準備金の額 [cite: 322][cite_start]\n金0円 [cite: 326]')
        [cite_start]doc.add_paragraph('Capital amount: [cite: 323][cite_start]\n100,000 yen [cite: 328]')
        [cite_start]doc.add_paragraph('Capital reserve: [cite: 324][cite_start]\n0 yen [cite: 329]')
        
        [cite_start]doc.add_paragraph('9. 払込取扱場所 [cite: 330][cite_start]\n東京都港区西麻布四丁目1番3号 [cite: 331][cite_start]\n株式会社三井住友銀行 [cite: 332][cite_start]\n麻布支店 [cite: 333]')
        [cite_start]doc.add_paragraph('Payment handling institution: [cite: 334][cite_start]\nSumitomo Mitsui Banking Corporation [cite: 335][cite_start]\nAzabu Branch [cite: 336][cite_start]\n4-1-3 Nishi-azabu, Minato-ku, Tokyo, Japan [cite: 337]')

        # --- PAGE 14 ---
        [cite_start]doc.add_paragraph('10. 会社法第31条第1項に基づく当会社成立前の定款の備置場所 [cite: 338]')
        [cite_start]doc.add_paragraph('Place to keep the Articles of Incorporation prior to the Company\'s incorporation pursuant to Article 31, Paragraph 1 of the Companies Act: [cite: 339]')
        [cite_start]doc.add_paragraph('東京都港区北青山二丁目9番5号スタジアムプレイス青山8階 当会社創立事務所 [cite: 340]')
        [cite_start]doc.add_paragraph('8th Floor STUDIUM PLACE AOYAMA, 2-9-5 Kita-aoyama, Minato-ku, Tokyo, Japan\nThe organization office of the Company [cite: 341]')

        [cite_start]doc.add_paragraph('11. 発起人は、当会社の設立に関し、報酬その他の特別の利益を受けないものとし、また、現物出資をしないものとする。 [cite: 342]')
        [cite_start]doc.add_paragraph('The Incorporator shall receive no remuneration or special interest in connection with the incorporation of the Company and shall make no investment in kind. [cite: 343]')

        [cite_start]doc.add_paragraph('12. 設立費用は、発起人が負担する。ただし、定款の認証の手数料、定款に係る印紙税、設立時発行株式と引換えにする金銭の払込みの取扱いをした銀行等に支払うべき手数料および報酬ならびに当会社の設立の登記の登録免許税は、当会社の負担とする。 [cite: 344, 345]')
        [cite_start]doc.add_paragraph('The Incorporator shall bear all the expenses for the incorporation; provided, however, that the Company shall bear the fee for notarization of the Articles of Incorporation, stamp duty for the Articles of Incorporation, fees and remuneration payable to the bank and the like for handling the payment which is in exchange for the shares to be issued at the incorporation and the registration and license tax for the incorporation. [cite: 345, 346]')

        [cite_start]doc.add_paragraph('13. 発起人は、発起人の決定に基づいて定款の作成、出資の履行その他会社の設立に関する一切の事務を執行する。 [cite: 347]')
        [cite_start]doc.add_paragraph('The Incorporator shall be solely responsible for all acts relating to the incorporation of the Company, including preparation of the Articles of Incorporation and execution of subscription. [cite: 348]')
        
        [cite_start]doc.add_paragraph('以上 End. [cite: 349]')
        
        [cite_start]doc.add_paragraph('以上の決定を証するため、本決定書を作成し、発起人が記名押印する。 [cite: 350]')
        [cite_start]doc.add_paragraph('IN WITNESS WHEREOF, this Statement of Incorporator has been prepared and the Incorporator has affixed his name and seal hereunto. [cite: 351]')
        
        [cite_start]doc.add_paragraph('平成28年12月7日 [cite: 353][cite_start]\nDecember, 2016 [cite: 354]')
        [cite_start]doc.add_paragraph('Adgorithmics株式会社 [cite: 356][cite_start]\nAdgorithmics KK [cite: 357]')
        [cite_start]doc.add_paragraph('発起人: [cite: 358][cite_start]\n東京都港区六本木六丁目12番2-1110号 [cite: 359][cite_start]\nキム・マイケル・サン・エン [cite: 360][cite_start]\n引受設立時発行株式数 1株 [cite: 361]')
        [cite_start]doc.add_paragraph('Incorporator: [cite: 362][cite_start]\nKim Michael Sang Eun [cite: 363][cite_start]\n6-12-2-1110 Roppongi, Minato-ku, Tokyo, Japan [cite: 365][cite_start]\nNumber of shares issued at incorporation to be allotted: 1 share [cite: 366]')
        [cite_start]doc.add_paragraph('[4字削除 4字加入] [cite: 367]')
        [cite_start]doc.add_paragraph('[Personal Seal] [cite: 368]')

        # --- PAGE 16 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('発起人決定書 [cite: 369]', level=1)
        [cite_start]doc.add_heading('STATEMENT OF INCORPORATOR [cite: 370]', level=1)
        [cite_start]doc.add_paragraph('平成28年12月8日、Adgorithmics株式会社(以下「当会社」という。)の発起人は、東京都港区北青山二丁目9番5号スタジアムプレイス青山8階当会社創立事務所において、下記事項を定めた。 [cite: 371]')
        [cite_start]doc.add_paragraph('On December 8, 2016, the Incorporator of Adgorithmics KK (the "Company"), determined the following matters at 8th Floor STUDIUM PLACE AOYAMA, 2-9-5 Kita-aoyama, Minato-ku, Tokyo, Japan, the organization office of the Company: [cite: 372]')
        [cite_start]doc.add_paragraph('記 Note [cite: 373]')
        
        [cite_start]doc.add_paragraph('1. 本店の所在場所 [cite: 374][cite_start]\n東京都港区北青山二丁目9番5号スタジアムプレイス青山8階 [cite: 375]')
        [cite_start]doc.add_paragraph('Location of the head office of the Company [cite: 376][cite_start]\n8th Floor STUDIUM PLACE AOYAMA, 2-9-5 Kita-aoyama, Minato-ku, Tokyo, Japan [cite: 377]')
        
        [cite_start]doc.add_paragraph('以上 End. [cite: 378]')
        
        [cite_start]doc.add_paragraph('以上の決定を証するため、本決定書を作成し、発起人が記名押印する。 [cite: 379]')
        [cite_start]doc.add_paragraph('IN WITNESS WHEREOF, this Statement of Incorporator has been prepared and the Incorporator has affixed his name and seal hereunto. [cite: 380]')
        
        [cite_start]doc.add_paragraph('平成28年12月8日 [cite: 381][cite_start]\nDecember 8, 2016 [cite: 382]')
        [cite_start]doc.add_paragraph('Adgorithmics株式会社 [cite: 383][cite_start]\nAdgorithmics KK [cite: 384]')
        [cite_start]doc.add_paragraph('発起人: [cite: 385][cite_start]\n東京都港区六本木六丁目12番2-1110号 [cite: 386][cite_start]\nキム・マイケル・サン・エン [cite: 387][cite_start]\n引受設立時発行株式数 1株 [cite: 388]')
        [cite_start]doc.add_paragraph('Incorporator : [cite: 389][cite_start]\nKim Michael Sang Eun [cite: 390][cite_start]\n6-12-2-1110 Roppongi, Minato-ku, Tokyo, Japan [cite: 391][cite_start]\nNumber of shares issued at incorporation to be allotted: 1 share [cite: 392]')
        [cite_start]doc.add_paragraph('[KIM] [cite: 393][cite_start]\n[Personal Seal] [cite: 394]')

        # --- PAGE 17 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('就任承諾書 [cite: 395]', level=1)
        [cite_start]doc.add_heading('ACCEPTANCE OF OFFICE [cite: 396]', level=1)
        [cite_start]doc.add_paragraph('私は、平成28年12月今日付け貴社定款の定めにより設立時取締役に選任されましたので、その就任を承諾いたします。 [cite: 397]')
        [cite_start]doc.add_paragraph('I have been elected as Director at incorporation of Adgorithmics KK (the "Company") by the Articles of Incorporation of the Company as of December, 2016, hereby accepts the office, effective as of this day. [cite: 398]')
        
        [cite_start]doc.add_paragraph('Adgorithmics株式会社 御中 [cite: 399][cite_start]\nTo: Adgorithmics KK [cite: 400]')
        
        [cite_start]doc.add_paragraph('平成28年12月7日 [cite: 402][cite_start]\nDecember 7, 2016 [cite: 403]')
        
        [cite_start]doc.add_paragraph('住所:東京都港区六本木六丁目12番2-1110号 [cite: 406][cite_start]\nAddress: 6-12-2-1110 Roppongi, Minato-ku, Tokyo, Japan [cite: 407]')
        [cite_start]doc.add_paragraph('氏名: キム・マイケル・サン・エン [cite: 408][cite_start]\nName: Kim Michael Sang Eun [cite: 409]')
        [cite_start]doc.add_paragraph('[KIM] [cite: 411][cite_start]\n[4字削除 千字加入。] [cite: 412, 413]')
        [cite_start]doc.add_paragraph('[Personal Seal] [cite: 414, 415]')

        # --- PAGE 18 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('証明書 [cite: 416]', level=1)
        [cite_start]doc.add_heading('CERTIFICATE [cite: 416]', level=1)
        [cite_start]doc.add_paragraph('当会社の設立時発行株式については、以下のとおり、全額の払込みがあったことを証明します。 [cite: 417]')
        [cite_start]doc.add_paragraph('I hereby certify that the subscription money for the shares issued at the incorporation has been paid in full as follows: [cite: 417]')
        
        add_table_from_data(doc, [
            [cite_start]['設立時発行株式数 [cite: 418][cite_start]', '1株 [cite: 418]'],
            [cite_start]['払込みを受けた金額 [cite: 418][cite_start]', '金100,000円 [cite: 418]']
        ])
        
        add_table_from_data(doc, [
            [cite_start]['Number of the share issued at the incorporation: [cite: 419][cite_start]', '1 share [cite: 419]'],
            [cite_start]['Amount of the subscription money for such shares: [cite: 419][cite_start]', '100,000 yen [cite: 419]']
        ])
        
        [cite_start]doc.add_paragraph('平成28年12月8日 [cite: 420][cite_start]\nDecember 8, 2016 [cite: 421]')
        [cite_start]doc.add_paragraph('Adgorithmics 株式会社 [cite: 422][cite_start]\nAdgorithmics KK [cite: 423]')
        [cite_start]doc.add_paragraph('設立時代表取締役: キム・マイケル・サン・エン [cite: 424, 425]')
        [cite_start]doc.add_paragraph('Representative Director at incorporation: Kim Michael Sang Eun [cite: 426]')
        [cite_start]doc.add_paragraph('[Company Seal] [cite: 427, 428]')

        # --- PAGE 19 ---
        doc.add_page_break()
        doc.add_heading('三井住友銀行 (SMBC) Bank Statement', level=2)
        [cite_start]doc.add_paragraph('三井住友銀行 [cite: 430][cite_start]\nSMBCダイレクト [cite: 432]')
        [cite_start]doc.add_paragraph('払 マイケル サン エンさまの残高・入出金明細は以下のとおりです。 [cite: 433]')
        [cite_start]doc.add_paragraph('平成28年12月3日 01:21現在 [cite: 434]')
        
        add_table_from_data(doc, [
            [cite_start]['照会口座 [cite: 435][cite_start]', '麻布支店 普通(総合) 535091 [cite: 435]'],
            [cite_start]['現在残高 [cite: 435][cite_start]', '183,974円 (お支払い可能残高:183,974円) [cite: 435]']
        ])
        
        [cite_start]doc.add_paragraph('平成28年11月1日から平成28年12月3日までの入出金明細 [cite: 436]')
        
        bank_data = [
            ["年月日", "お引出し", "お預入れ", "お取り扱い内容", "残高"],
            ["H28.11.2", "199,000円:", "", "カード エスエムビーシーシンタク", "264,226円"],
            ["H28.11.4", "24,862 円", "", "セブンCSカード", "239,364円"],
            ["H28.11.4", "", "705,893円", "外国関係 ヒシムケソウキン49379927", "945,257円"],
            ["H28.11.4", "566,909円", "", "ローンご返済", "378,348円"],
            ["H28.11.8", "199,000円", "", "カード ローソン234391", "179,348円"],
            ["H28.11.8", "216円", "", "カード手数料", "179,132円"],
            ["H28.11.14", "150,000円", "", "カード エスエムビーシーシンタク", "29,132円"],
            ["H28.11.24", "", "994,000円", "外国関係 ヒシムケソウキン49390142", "1,023,132円"],
            ["H28.11.25", "", "449,830円", "振込サービス アドゴリズミクス エルエルシー", "1,472,962円"],
            ["H28.11.27", "199,000円", "", "カード ユウチョ 017240", "1,273,962円"],
            ["H28.11.27", "216円", "", "カード手数料", "1,273,746円"],
            ["H28.11.28", "565,727円", "", "ローンご返済", "708,019円"],
            ["H28.11.28", "463,648円", "", "エムアイカード", "244,371円"],
            ["H28.11.28", "6,337円", "", "DF.ロッポンギヒルズ", "238,034円"],
            ["H28.11.28", "69,960円", "", "DF、カンリヒトウ", "168,074円"],
            ["H28.11.28", "17,000円", "", "DF.BFA", "151,074円"],
            ["H28.11.28", "45,700円", "", "オリコ", "105,374円"],
            ["H28.11.30", "21,400円", "", "ミナトク ケンポトコウキン", "83,974円"],
            ["H28.12.2", "", "100,000円", "振込サービス アドゴリズミクス エルエルシー", "183,974円"],
            ["合計金額", "2,528,975円!", "2,249,723円", "", ""]
        ]
        # Adding data from cite: 437
        add_table_from_data(doc, bank_data)
        doc.add_paragraph('Copyright © 2016 Sumitomo Mitsui Banking Corporation. [cite_start]All Rights Reserved. [cite: 439]')

        # --- PAGE 21 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('調査報告書 [cite: 440]', level=1)
        [cite_start]doc.add_heading('INVESTIGATION REPORT [cite: 441]', level=1)
        [cite_start]doc.add_paragraph('私は設立時取締役に選任されたので、会社法第46条の規定に基づいて調査をした結果を、次のとおり報告する。 [cite: 442]')
        [cite_start]doc.add_paragraph('Being appointed as a Director at incorporation, we report the result of our investigation made under Article 46 of the Companies Act as follows: [cite: 443]')
        
        [cite_start]doc.add_paragraph().add_run('調査事項 [cite: 444]').bold = True
        [cite_start]doc.add_paragraph().add_run('ITEMS OF INVESTIGATION [cite: 445]').bold = True
        
        [cite_start]doc.add_paragraph('1. 平成28年12月2日までに出資の履行が完了していることは、発起人の口座の預金通帳の記録により認められる。 [cite: 446]')
        [cite_start]doc.add_paragraph('We confirm that the execution of the subscription by the Incorporator has been completed by December 2, 2016, according to the record of the Incorporator\'s bank book. [cite: 447]')
        
        [cite_start]doc.add_paragraph('2. 上記事項のほか、株式会社の設立の手続が法令または定款に違反していないことを認める。 [cite: 448]')
        [cite_start]doc.add_paragraph('We confirm that the procedures concerning the incorporation, besides the above item, do not violate any laws and regulations or the Articles of Incorporation. [cite: 449]')
        
        [cite_start]doc.add_paragraph('平成28年12月8日 [cite: 450][cite_start]\nDecember 8, 2016 [cite: 451]')
        [cite_start]doc.add_paragraph('Adgorithmics株式会社 [cite: 452][cite_start]\nAdgorithmics KK [cite: 453]')
        [cite_start]doc.add_paragraph('設立時取締役 [cite: 454][cite_start]\nDirector at incorporation [cite: 455]')
        [cite_start]doc.add_paragraph('キム・マイケル・サン・エン [cite: 456][cite_start]\nKim Michael Sang Eun [cite: 456]')
        [cite_start]doc.add_paragraph('[Personal Seal] [cite: 457]')

        # --- PAGE 22 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('印鑑登録証明書 [cite: 459]', level=1)
        doc.add_heading('Certificate of Seal Registration', level=2)
        [cite_start]doc.add_paragraph('Minato City [cite: 460]')
        
        seal_data = [
            [cite_start]['氏名 [cite: 461][cite_start]', 'KIM MICHAEL SANG EUN [cite: 461][cite_start]', '生年月日 [cite: 461][cite_start]', '1963年 9月25日 [cite: 461]'],
            [cite_start]['通称 [cite: 461]', '', '', ''],
            [cite_start]['住所 [cite: 461][cite_start]', '東京都港区六本木6丁目12番2-1110号 [cite: 461]', '', ''],
            [cite_start]['備考 [cite: 461]', '', '', '']
        ]
        add_table_from_data(doc, seal_data)
        
        [cite_start]doc.add_paragraph('印影 [cite: 462][cite_start]\n[KIM] [cite: 463]')
        [cite_start]doc.add_paragraph('Minato City [cite: 464]')
        [cite_start]doc.add_paragraph('この写しは、登録された印影と相違ないことを証明します。 [cite: 465]')
        [cite_start]doc.add_paragraph('平成28年11月9日 [cite: 466]')
        [cite_start]doc.add_paragraph('港赤印証 第660441-01号 [cite: 467]')
        [cite_start]doc.add_paragraph('東京都港区長 [cite: 468][cite_start]\n武井雅昭 [cite: 469]')
        [cite_start]doc.add_paragraph('[Anti-forgery text and stamps] [cite: 470, 471, 472, 473, 474, 475, 476]')

        # --- PAGE 23 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('委任状 [cite: 477]', level=1)
        [cite_start]doc.add_heading('POWER OF ATTORNEY FOR NOTARIZATION [cite: 478]', level=1)
        [cite_start]doc.add_paragraph('住所: [cite: 479][cite_start]\nADDRESS: [cite: 480]')
        [cite_start]doc.add_paragraph('氏名: [cite: 481][cite_start]\nNAME: [cite: 482]')
        [cite_start]doc.add_paragraph('私は、上記の者を代理人と定め、以下の権限を委託する。 [cite: 483]')
        [cite_start]doc.add_paragraph('KNOW ALL MEN BY THESE PRESENTS, that I do hereby make, constitute and appoint the above mentioned person as my true and lawful attorney-in-fact in respect of the following items: [cite: 484, 485]')
        
        [cite_start]doc.add_paragraph('1 東京法務局所属公証人に対して、Adgorithmics株式会社の設立のために原始定款の認証を請求し原本を受領する件 [cite: 486]')
        [cite_start]doc.add_paragraph('但し、下記発起人は定款になした記名押印を自認する。 [cite: 488]')
        [cite_start]doc.add_paragraph('To apply to a notary public of the Tokyo Legal Affairs Bureau, for notarization of the original Articles of Incorporation of Adgorithmics KK and to receive an original copy of the notarized, original Articles of Incorporation, provided that the Incorporator listed below shall acknowledge his own signature affixed in the Articles of Incorporation; and [cite: 489, 490]')
        
        [cite_start]doc.add_paragraph('定款謄本の交付請求および受領に関する件 [cite: 490]')
        [cite_start]doc.add_paragraph('To apply for and receive the certified copy of the notarized, original Articles of Incorporation. [cite: 491]')
        
        [cite_start]doc.add_paragraph('平成28年12月8日 [cite: 492][cite_start]\nDecember 8, 2016 [cite: 493]')
        [cite_start]doc.add_paragraph('会社名: Adgorithmics株式会社 [cite: 494, 495]')
        [cite_start]doc.add_paragraph('Name of the Company: Adgorithmics KK [cite: 496]')
        [cite_start]doc.add_paragraph('発起人: [cite: 497][cite_start]\n東京都港区六本木六丁目12番2-1110号 [cite: 498][cite_start]\nキム・マイケル・サン・エン [cite: 499][cite_start]\n引受設立時発行株式数 1株 [cite: 500]')
        [cite_start]doc.add_paragraph('Incorporator: [cite: 501][cite_start]\nKim Michael Sang Eun [cite: 502][cite_start]\n6-12-2-1110 Roppongi, Minato-ku, Tokyo, Japan [cite: 503][cite_start]\nNumber of shares issued at incorporation to be allotted: 1 share [cite: 504]')
        [cite_start]doc.add_paragraph('[KIM] [cite: 505][cite_start]\n[Personal Seal] [cite: 506]')

        # --- PAGE 24 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('委任状 [cite: 507]', level=1)
        [cite_start]doc.add_heading('POWER OF ATTORNEY [cite: 508]', level=1)
        [cite_start]doc.add_paragraph('私は、東京都港区六本木六丁目10番1号六本木ヒルズ森タワー23階TMI総合法律事務所弁護士平野正弥を代理人と定め、下記の権限を委任します。 [cite: 509]')
        [cite_start]doc.add_paragraph('KNOW ALL MEN BY THESE PRESENTS, that I do hereby make, constitute and appoint Mr. Masaya Hirano, Attorney at-law, TMI Associates, 23rd Floor, Roppongi Hills Mori Tower, 6-10-1 Roppongi, Minato-ku, Tokyo, as our true and lawful attorney-in-fact for the following items: [cite: 510, 511]')
        
        [cite_start]doc.add_paragraph('記 [cite: 512][cite_start]\nNote [cite: 513]')
        [cite_start]doc.add_paragraph('1. 当会社の設立の登記申請に関する一切の件 [cite: 514]\n1. [cite_start]To file an application for registration of the incorporation of the Company; [cite: 515, 516]')
        [cite_start]doc.add_paragraph('1. 原本還付請求および受領の件 [cite: 518]\n1. [cite_start]To apply for and receive an original set of registration documents; [cite: 519, 520]')
        [cite_start]doc.add_paragraph('1. 申請補正を要する場合、補正または取下げをなす件 [cite: 522]\n1. [cite_start]To modify and revoke the application documents, if necessary; and [cite: 523, 524]')
        [cite_start]doc.add_paragraph('1. 上記に関するその他一切の件 [cite: 526]\n1. [cite_start]Any other transactions related to any of the above items. [cite: 527, 528]')
        
        [cite_start]doc.add_paragraph('以上 [cite: 529][cite_start]\nEnd. [cite: 530]')
        
        [cite_start]doc.add_paragraph('平成28年12月8日 [cite: 531][cite_start]\nDecember 8, 2016 [cite: 532]')
        [cite_start]doc.add_paragraph('東京都港区北青山二丁目9番5号スタジアムプレイス青山8階 [cite: 533][cite_start]\nAdgorithmics 株式会社 [cite: 533][cite_start]\n代表取締役 キム・マイケル・サン・エン [cite: 534]')
        [cite_start]doc.add_paragraph('8th Floor STUDIUM PLACE AOYAMA, 2-9-5 Kita-aoyama, Minato-ku, Tokyo, Japan [cite: 535, 536][cite_start]\nAdgorithmics KK [cite: 536][cite_start]\nRepresentative Director: Kim Michael Sang Eun [cite: 537]')
        [cite_start]doc.add_paragraph('[Company seal] [cite: 538]')

        # --- PAGE 25 ---
        doc.add_page_break()
        [cite_start]doc.add_heading('印鑑(改印)届書 [cite: 539]', level=1)
        doc.add_paragraph('Seal (Change) Registration Form', 'Heading 2')
        [cite_start]doc.add_paragraph('※太枠の中に書いてください。 [cite: 540]')
        
        [cite_start]doc.add_paragraph().add_run('商号・名称: [cite: 544]').bold = True
        [cite_start]doc.add_paragraph('Adgorithmics株式会社 [cite: 545]')
        
        [cite_start]doc.add_paragraph().add_run('本店・主たる事務所: [cite: 546]').bold = True
        [cite_start]doc.add_paragraph('東京都港区北青山二丁目9番5号 スタジアムプレイス青山8階 [cite: 547]')

        [cite_start]doc.add_paragraph().add_run('印鑑提出者 [cite: 549]').bold = True
        [cite_start]doc.add_paragraph('資格: [cite: 550] [cite_start]代表取締役 [cite: 548]')
        [cite_start]doc.add_paragraph('氏名: [cite: 553] [cite_start]キム・マイケル・サン・エン [cite: 554]')
        [cite_start]doc.add_paragraph('生年月日: [cite: 555] [cite_start]1963年9月25日生 [cite: 556]')

        [cite_start]doc.add_paragraph().add_run('印鑑カード: [cite: 558]').bold = True
        [cite_start]doc.add_paragraph('☑ 印鑑カードは引き継がない。 [cite: 558]')
        
        [cite_start]doc.add_paragraph().add_run('届出人: [cite: 563]').bold = True
        [cite_start]doc.add_paragraph('☑ 代理人 [cite: 563]')
        [cite_start]doc.add_paragraph('住所: [cite: 564] [cite_start]東京都港区六本木六丁目10番1号六本木ヒルズ森タワー23階 TMI総合法律事務所 [cite: 566]')
        [cite_start]doc.add_paragraph('氏名: [cite: 567] [cite_start]弁護士 平野 正弥 [cite: 567]')

        [cite_start]doc.add_paragraph().add_run('委任状 [cite: 569]').bold = True
        [cite_start]doc.add_paragraph('私は、(住所) 東京都港区六本木六丁目10番1号六本木ヒルズ森タワー23階 TMI総合法律事務所 [cite: 571, 572]')
        [cite_start]doc.add_paragraph('(氏名) 弁護士平野 正弥 [cite: 573]')
        [cite_start]doc.add_paragraph('を代理人と定め、印鑑 (改印) の届出の権限を委任します。 [cite: 574]')
        [cite_start]doc.add_paragraph('平成28年12月8日 [cite: 575]')
        [cite_start]doc.add_paragraph('住所 東京都港区六本木六丁目12番2-1110号 [cite: 576]')
        [cite_start]doc.add_paragraph('氏名・キム・マイケル・サン・エン [cite: 577]')
        [cite_start]doc.add_paragraph('[Seal registered with municipality] [cite: 579, 578]')
        
        [cite_start]doc.add_paragraph('☑ 市区町村長作成の印鑑証明書は、登記申請書に添付のものを援用する。(注4) [cite: 581]')
        
        # Save the document
        output_filename = 'Adgorithmics_Incorporation_Documents.docx'
        doc.save(output_filename)
        print(f"Successfully created '{output_filename}'")

    except ImportError:
        print("Error: The 'python-docx' library is not installed.")
        print("Please install it by running: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"An error occurred: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()